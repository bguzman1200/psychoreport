import streamlit as st
import pandas as pd
from fuzzywuzzy import fuzz, process
from docx import Document
from docx.shared import Pt
import io
from openai import OpenAI

# -----------------------------------------------------------------------------
# CONSTANTS & CONFIGURATION
# -----------------------------------------------------------------------------
st.set_page_config(page_title="PsychoReport.ai", page_icon="🧠", layout="wide")

# Domain Keywords for Fuzzy Matching
DOMAINS = {
    "Behavioral Regulation": [
        "Blurts", "Interrupts", "Rushes", "Waiting", "Overreacts", 
        "Calm", "Transition", "Stuck", "Rigid"
    ],
    "Cognitive Regulation": [
        "Distracted", "Redirection", "Zoning", "Track", "Multi-step", 
        "Forgets", "Recall", "Prompting", "Procrastinates", "Check work"
    ],
    "Organization": [
        "Desk", "Messy", "Materials", "Lose items", "Break down", 
        "Overwhelmed", "Underestimates", "Time"
    ]
}

# -----------------------------------------------------------------------------
# HELPER FUNCTIONS
# -----------------------------------------------------------------------------

def score_text_to_int(value):
    """Converts Likert scale text (e.g., '1 (Never/Rarely)') to an integer."""
    if pd.isna(value):
        return None
    try:
        str_val = str(value).strip()
        if str_val and str_val[0].isdigit():
            return int(str_val[0])
        return None
    except:
        return None

def clean_and_score(df, rater_name):
    """Analyzes a raw dataframe using fuzzy matching to map columns to domains."""
    scores = {d: [] for d in DOMAINS.keys()}
    notes = []

    for col in df.columns:
        col_str = str(col)
        
        # 1. Check for Qualitative Data
        if any(x in col_str.lower() for x in ["note", "comment", "example", "elaborate"]):
            valid_notes = df[col].dropna().astype(str).tolist()
            if valid_notes:
                notes.extend([f"- {n}" for n in valid_notes if len(n) > 3])
            continue

        # 2. Fuzzy Match Column Header to Domains
        best_domain = None
        best_score = 0
        
        for domain, keywords in DOMAINS.items():
            matches = process.extract(col_str, keywords, scorer=fuzz.partial_ratio, limit=1)
            if matches:
                keyword, score = matches[0][0], matches[0][1]
                if score > 80: 
                    if score > best_score:
                        best_score = score
                        best_domain = domain
        
        # 3. Process Scores
        if best_domain:
            numeric_values = df[col].apply(score_text_to_int).dropna()
            scores[best_domain].extend(numeric_values.tolist())

    # Calculate Averages
    final_scores = {}
    for domain, val_list in scores.items():
        if val_list:
            final_scores[domain] = round(sum(val_list) / len(val_list), 2)
        else:
            final_scores[domain] = "N/A"
            
    return final_scores, "\n".join(notes)

def create_docx(report_text):
    """Generates a professional Word Document from the GPT report text."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    doc.add_heading('Psychoeducational Evaluation Report', 0)
    
    for line in report_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        if line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), level=1)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            clean_line = line.replace('**', '').replace('__', '')
            doc.add_paragraph(clean_line)
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------------------------------------------------------
# MAIN APP UI
# -----------------------------------------------------------------------------

with st.sidebar:
    st.title("Settings")
    api_key = st.text_input("Enter your OpenAI API Key", type="password")
    st.markdown("[Don't have a key? Get one at platform.openai.com](https://platform.openai.com)")
    st.divider()
    st.info("🔒 Privacy Mode: Data is processed in-memory and wiped on refresh. No data is stored.")

st.title("🧠 PsychoReport.ai")
st.markdown("### Automated Psychoeducational Reporting Tool")

# --- FIXED: Added Student Name Input to prevent NameError ---
student_name = st.text_input("Student Name (for Report Header)", placeholder="e.g. Alex M.")
# -----------------------------------------------------------

col1, col2, col3 = st.columns(3)
with col1:
    teacher_file = st.file_uploader("Teacher Data (CSV)", type="csv")
with col2:
    parent_file = st.file_uploader("Parent Data (CSV)", type="csv")
with col3:
    observer_file = st.file_uploader("Observer Data (CSV)", type="csv")

# -----------------------------------------------------------------------------
# LOGIC ENGINE
# -----------------------------------------------------------------------------

if st.button("Generate Report", type="primary"):
    if not api_key:
        st.error("⚠️ Authentication Error: Please provide an OpenAI API Key in the sidebar.")
        st.stop()
        
    if not (teacher_file or parent_file or observer_file):
        st.error("⚠️ Missing Data: Please upload at least one CSV file.")
        st.stop()
        
    if not student_name:
        st.warning("⚠️ Please enter a Student Name to generate the report.")
        st.stop()

    with st.spinner(f"Analyzing data for {student_name} and generating clinical report..."):
        try:
            data_context = f"Student Name: {student_name}\n"
            
            # Process Files
            if teacher_file:
                df_t = pd.read_csv(teacher_file)
                scores_t, notes_t = clean_and_score(df_t, "Teacher")
                data_context += f"\n--- TEACHER DATA ---\nScores (1-4): {scores_t}\nNotes:\n{notes_t}"
            else:
                data_context += "\n--- TEACHER DATA ---\n(Data not available)"

            if parent_file:
                df_p = pd.read_csv(parent_file)
                scores_p, notes_p = clean_and_score(df_p, "Parent")
                data_context += f"\n--- PARENT DATA ---\nScores (1-4): {scores_p}\nNotes:\n{notes_p}"
            else:
                data_context += "\n--- PARENT DATA ---\n(Data not available)"

            if observer_file:
                df_o = pd.read_csv(observer_file)
                scores_o, notes_o = clean_and_score(df_o, "Observer")
                data_context += f"\n--- OBSERVER DATA ---\nScores (1-4): {scores_o}\nNotes:\n{notes_o}"
            else:
                data_context += "\n--- OBSERVER DATA ---\n(Data not available)"

            # Prompt Construction
            system_prompt = (
                "You are a Licensed Educational Psychologist writing a formal psychoeducational evaluation. "
                "Tone: Clinical, Objective, Professional. "
                "Do not invent scores if they are missing (marked N/A)."
            )
            
            user_prompt = f"""
            Write a comprehensive evaluation report for: {student_name}.
            
            INPUT DATA:
            {data_context}
            
            REQUIRED REPORT STRUCTURE (Use Markdown Headers):
            
            # Executive Functioning Profile: {student_name}
            
            # Cross-Contextual Analysis
            Compare findings across settings (Home vs School). Highlight discrepancies.
            
            # Domain Analysis
            ## Behavioral Regulation
            Interpret the scores (1=Low Concern, 4=High Concern).
            ## Cognitive Regulation
            Interpret the scores.
            ## Organization & Planning
            Interpret the scores.
            
            # Recommendations
            ## School-Based Interventions
            Provide 3 specific strategies.
            ## Home-Based Interventions
            Provide 3 specific strategies.
            """

            client = OpenAI(api_key=api_key)
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7
            )
            
            report_content = response.choices[0].message.content
            
            st.subheader("Generated Report")
            st.text_area("Preview", value=report_content, height=600)
            
            docx_file = create_docx(report_content)
            
            st.download_button(
                label="📥 Download Report as Word Doc (.docx)",
                data=docx_file,
                file_name=f"PsychoReport_{student_name.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            st.error(f"An error occurred: {e}")
